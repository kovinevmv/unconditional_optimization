import os

from PyQt5 import QtCore
from docx import Document
from docx.shared import Inches

import resources


class Doc:
    def __init__(self, image, table, method_name):
        self.image_stream = image
        self.data = table
        self.method = method_name

    def save(self, name):
        self.name = name

        default_path = os.path.join(os.getenv("TEMP"), 'default.docx')
        QtCore.QFile.copy(':/help_docs/default.docx', default_path)

        self._save(default_path)
        try:
            os.remove(default_path)
        except:
            pass

    def _save(self, default_path):
        self.document = Document(default_path)

        self.document.add_heading(self.method, 0)
        self.document.add_heading('Рисунок', level=1)
        self.document.add_picture(self.image_stream, width=Inches(6))

        self.document.add_page_break()
        self.document.add_heading('Таблица', level=1)

        headers = list(self.data[0].keys())
        table = self.document.add_table(rows=len(self.data) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        for i, head in enumerate(headers):
            hdr_cells[i].text = head

        for i, d in enumerate(self.data):
            for j, (_, value) in enumerate(d.items()):
                table.cell(i + 1, j).text = str(value)

        self.document.save(self.name)


if __name__ == '__main__':
    pass
